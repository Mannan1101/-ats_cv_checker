"""Export Agent.

Renders the pipeline's results to disk: a Markdown ATS report, a PDF
version of the same report, an ATS-optimized DOCX resume, and a DOCX cover
letter. Pure Python/templating -- no LLM calls, so output is fast and
deterministic.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from models.schemas import (
    ATSAnalysisReport,
    CoverLetter,
    ImprovedResume,
    JobRequirements,
    ParsedResume,
)

_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATES_DIR)),
    autoescape=select_autoescape(enabled_extensions=()),
    trim_blocks=True,
    lstrip_blocks=True,
)


def _safe_stub(name: str | None, fallback: str) -> str:
    base = (name or fallback).strip().lower().replace(" ", "_")
    return "".join(c for c in base if c.isalnum() or c == "_") or fallback


def render_markdown_report(
    resume: ParsedResume,
    jd: JobRequirements,
    report: ATSAnalysisReport,
    improved: ImprovedResume,
) -> str:
    template = _env.get_template("report_template.md.j2")
    return template.render(
        resume=resume,
        jd=jd,
        report=report,
        improved=improved,
        generated_at=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
    )


def export_markdown_report(
    resume: ParsedResume,
    jd: JobRequirements,
    report: ATSAnalysisReport,
    improved: ImprovedResume,
    output_dir: str | Path,
) -> Path:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    stub = _safe_stub(resume.contact_info.full_name, "candidate")
    path = output_dir / f"{stub}_ats_report.md"
    path.write_text(render_markdown_report(resume, jd, report, improved), encoding="utf-8")
    return path


def export_pdf_report(
    resume: ParsedResume,
    jd: JobRequirements,
    report: ATSAnalysisReport,
    improved: ImprovedResume,
    output_dir: str | Path,
) -> Path:
    """Render the same content as the Markdown report, as a PDF."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    stub = _safe_stub(resume.contact_info.full_name, "candidate")
    path = output_dir / f"{stub}_ats_report.pdf"

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6)
    body = styles["BodyText"]
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=14, bulletIndent=2)

    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.8 * inch, bottomMargin=0.8 * inch,
    )
    story = [
        Paragraph("ATS Compatibility Report", h1),
        Paragraph(f"Candidate: {resume.contact_info.full_name or 'N/A'}", body),
        Paragraph(f"Target Role: {jd.job_title or 'N/A'}" + (f" at {jd.company_name}" if jd.company_name else ""), body),
        Spacer(1, 12),
        Paragraph(f"<b>ATS Score: {report.ats_score} / 100</b>", h2),
        Paragraph(report.summary or "", body),
        Paragraph("Score Breakdown", h2),
    ]
    for label, value in [
        ("Keyword Match (30%)", report.match_breakdown.keyword_match),
        ("Skills Match (25%)", report.match_breakdown.skills_match),
        ("Experience Match (20%)", report.match_breakdown.experience_match),
        ("Project Match (10%)", report.match_breakdown.project_match),
        ("Education Match (10%)", report.match_breakdown.education_match),
        ("Formatting (5%)", report.match_breakdown.formatting_score),
    ]:
        story.append(Paragraph(f"{label}: {value}%", bullet))

    story.append(Paragraph("Strengths", h2))
    for s in report.strengths:
        story.append(Paragraph(f"&bull; {s}", bullet))

    story.append(Paragraph("Weaknesses", h2))
    for w in report.weaknesses:
        story.append(Paragraph(f"&bull; {w}", bullet))

    story.append(Paragraph("Missing Keywords", h2))
    story.append(Paragraph(", ".join(report.missing_keywords) or "None", body))

    story.append(Paragraph("Missing Skills", h2))
    story.append(Paragraph(", ".join(report.missing_skills) or "None", body))

    story.append(Paragraph("Improved Professional Summary", h2))
    story.append(Paragraph(improved.improved_summary, body))

    story.append(Paragraph("Key Achievements", h2))
    for a in improved.achievements:
        story.append(Paragraph(f"&bull; {a}", bullet))

    doc.build(story)
    return path


def export_optimized_resume_docx(
    resume: ParsedResume, improved: ImprovedResume, output_dir: str | Path
) -> Path:
    """Build an ATS-friendly DOCX resume using the improved content."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    stub = _safe_stub(resume.contact_info.full_name, "candidate")
    path = output_dir / f"{stub}_optimized_resume.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    contact = resume.contact_info
    doc.add_heading(contact.full_name or "Candidate", level=0)
    contact_line = " | ".join(
        v for v in [contact.email, contact.phone, contact.location, contact.linkedin, contact.github]
        if v
    )
    if contact_line:
        doc.add_paragraph(contact_line)

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(improved.improved_summary)

    if improved.improved_skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(improved.improved_skills))

    if improved.improved_experience:
        doc.add_heading("Experience", level=1)
        # match improved bullets back to their original job entry (order-based)
        for exp, exp_source in zip(improved.improved_experience, resume.work_experience):
            title_line = f"{exp.job_title} — {exp.company}"
            dates = " to ".join(filter(None, [exp_source.start_date, "Present" if exp_source.is_current else exp_source.end_date]))
            p = doc.add_paragraph()
            p.add_run(title_line).bold = True
            if dates:
                p.add_run(f"  ({dates})")
            for b in exp.improved_bullets:
                doc.add_paragraph(b.improved, style="List Bullet")

    if improved.improved_projects:
        doc.add_heading("Projects", level=1)
        for proj in improved.improved_projects:
            p = doc.add_paragraph()
            p.add_run(proj.name).bold = True
            for b in proj.improved_bullets:
                doc.add_paragraph(b.improved, style="List Bullet")

    if resume.education:
        doc.add_heading("Education", level=1)
        for edu in resume.education:
            line = f"{edu.degree}, {edu.institution}"
            if edu.field_of_study:
                line += f" — {edu.field_of_study}"
            doc.add_paragraph(line, style="List Bullet")

    if resume.certifications:
        doc.add_heading("Certifications", level=1)
        for cert in resume.certifications:
            doc.add_paragraph(f"{cert.name}" + (f" — {cert.issuer}" if cert.issuer else ""), style="List Bullet")

    if improved.achievements:
        doc.add_heading("Key Achievements", level=1)
        for a in improved.achievements:
            doc.add_paragraph(a, style="List Bullet")

    doc.save(str(path))
    return path


def export_cover_letter_docx(
    cover_letter: CoverLetter, resume: ParsedResume, output_dir: str | Path
) -> Path:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    stub = _safe_stub(resume.contact_info.full_name, "candidate")
    path = output_dir / f"{stub}_cover_letter.docx"

    doc = Document()
    for paragraph in cover_letter.full_text.split("\n\n"):
        doc.add_paragraph(paragraph.strip())
    doc.save(str(path))
    return path


def export_all(
    resume: ParsedResume,
    jd: JobRequirements,
    report: ATSAnalysisReport,
    improved: ImprovedResume,
    output_dir: str | Path,
    cover_letter: CoverLetter | None = None,
) -> dict[str, str]:
    """Run every exporter and return {kind: path} for the coordinator/API."""
    exported = {
        "markdown_report": str(export_markdown_report(resume, jd, report, improved, output_dir)),
        "pdf_report": str(export_pdf_report(resume, jd, report, improved, output_dir)),
        "optimized_resume_docx": str(export_optimized_resume_docx(resume, improved, output_dir)),
    }
    if cover_letter is not None:
        exported["cover_letter_docx"] = str(
            export_cover_letter_docx(cover_letter, resume, output_dir)
        )
    return exported

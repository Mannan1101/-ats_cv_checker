from pathlib import Path

import pytest
from docx import Document

from services.docx_reader import extract_text_from_docx, read_resume_text


def test_extract_text_from_docx(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Jordan Ellis")
    doc.add_paragraph("Backend Engineer with 5 years of experience.")
    path = tmp_path / "resume.docx"
    doc.save(str(path))

    text = extract_text_from_docx(path)
    assert "Jordan Ellis" in text
    assert "Backend Engineer" in text


def test_read_resume_text_txt(tmp_path: Path):
    path = tmp_path / "resume.txt"
    path.write_text("Plain text resume content", encoding="utf-8")
    assert read_resume_text(path) == "Plain text resume content"


def test_read_resume_text_unsupported_extension(tmp_path: Path):
    path = tmp_path / "resume.xyz"
    path.write_text("data", encoding="utf-8")
    with pytest.raises(ValueError):
        read_resume_text(path)


def test_read_resume_text_missing_file(tmp_path: Path):
    with pytest.raises(FileNotFoundError):
        read_resume_text(tmp_path / "does_not_exist.pdf")

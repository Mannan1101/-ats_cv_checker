"""DOCX text extraction and writing helpers."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_text_from_docx(path: str | Path) -> str:
    """Extract plain text (paragraphs + table cells) from a DOCX file."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    document = Document(str(path))
    chunks: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)

    return "\n".join(chunks).strip()


def read_resume_text(path: str | Path) -> str:
    """Dispatch to the correct extractor based on file extension."""
    from services.pdf_reader import extract_text_from_pdf

    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported resume file type: {suffix}")
